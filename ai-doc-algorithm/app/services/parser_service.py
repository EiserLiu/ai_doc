import io
import logging

import fitz  # PyMuPDF
from docx import Document
import openpyxl

logger = logging.getLogger(__name__)


def parse_file(file_bytes: bytes, file_ext: str) -> str:
    if file_ext == "pdf":
        return parse_pdf(file_bytes)
    elif file_ext == "docx":
        return parse_docx(file_bytes)
    elif file_ext == "xlsx":
        return parse_xlsx(file_bytes)
    else:
        raise ValueError(f"不支持的文件类型: .{file_ext}")


def parse_pdf(file_bytes: bytes) -> str:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    texts = []
    for page in doc:
        text = page.get_text("text")
        if text.strip():
            texts.append(text.strip())
    doc.close()
    return "\n\n".join(texts)


def parse_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    texts = []

    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                texts.append(" | ".join(cells))

    return "\n\n".join(texts)


def parse_xlsx(file_bytes: bytes) -> str:
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    texts = []
    for sheet in wb.worksheets:
        sheet_texts = []
        for row in sheet.iter_rows(values_only=True):
            cells = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
            if cells:
                sheet_texts.append(" | ".join(cells))
        if sheet_texts:
            texts.append(f"[Sheet: {sheet.title}]\n" + "\n".join(sheet_texts))
    wb.close()
    return "\n\n".join(texts)
